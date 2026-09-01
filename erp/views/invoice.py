"""
erp/views/invoice.py
GST Tax Invoice generation — CTO Logistics & Infra format.

Run this SQL once in Supabase before using:
    CREATE TABLE invoices (
        id            UUID DEFAULT gen_random_uuid() PRIMARY KEY,
        invoice_number TEXT UNIQUE NOT NULL,
        work_order_id  UUID,
        invoice_date   DATE NOT NULL,
        customer_id    UUID,
        site_id        UUID,
        tax_type       TEXT DEFAULT 'CGST/SGST',
        line_items     JSONB,
        subtotal       NUMERIC(14,2) DEFAULT 0,
        tax_amount     NUMERIC(14,2) DEFAULT 0,
        round_off      NUMERIC(6,2)  DEFAULT 0,
        grand_total    NUMERIC(14,2) DEFAULT 0,
        status         TEXT DEFAULT 'Draft',
        notes          TEXT,
        created_at     TIMESTAMPTZ DEFAULT NOW()
    );
"""
from __future__ import annotations

import json
import calendar
from datetime import date, datetime

import streamlit as st

from ..supabase_client import SupabaseClient
from erp.views._documents import render_document_panel

# ── Fixed company / bank constants ────────────────────────────────────────────
_CO = {
    "name":       "CTO LOGISTICS & INFRA",
    "gstin":      "27AASFC8920H1Z1",
    "state":      "Maharashtra",
    "state_code": "27",
    "email":      "ctologinfra@gmail.com",
    "tel":        "022-4215 6953",
    "addr1":      "B-202, Second, Steel Chambers,",
    "addr2":      "Steel Market Road, Plot No. 514,",
    "addr3":      "Varsha Cranes Pvt. Ltd. Kalamboli,",
    "addr4":      "Navi Mumbai - 410218",
}
_BANK = {
    "holder":  "CTO Logistics & Infra",
    "bank":    "ICICI Bank",
    "account": "109805002451",
    "ifsc":    "ICIC0001098",
}


# ── Utility helpers ────────────────────────────────────────────────────────────

def _fy_str(d: date | None = None) -> str:
    d = d or date.today()
    y = d.year if d.month >= 4 else d.year - 1
    return f"{str(y)[2:]}-{str(y + 1)[2:]}"


_ONES = [
    "", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
    "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
    "Seventeen", "Eighteen", "Nineteen",
]
_TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty",
         "Sixty", "Seventy", "Eighty", "Ninety"]


def _w2(n: int) -> str:
    return _ONES[n] if n < 20 else _TENS[n // 10] + (" " + _ONES[n % 10] if n % 10 else "")


def _w3(n: int) -> str:
    if n < 100:
        return _w2(n)
    return _ONES[n // 100] + " Hundred" + (" " + _w2(n % 100) if n % 100 else "")


def _num_to_words_inr(amount: float) -> str:
    n = int(round(max(0.0, amount)))
    if n == 0:
        return "Zero Rupees Only"
    parts: list[str] = []
    if n >= 1_00_00_000:
        parts.append(_w3(n // 1_00_00_000) + " Crore"); n %= 1_00_00_000
    if n >= 1_00_000:
        parts.append(_w3(n // 1_00_000) + " Lakh"); n %= 1_00_000
    if n >= 1_000:
        parts.append(_w3(n // 1_000) + " Thousand"); n %= 1_000
    if n >= 100:
        parts.append(_ONES[n // 100] + " Hundred"); n %= 100
    if n:
        parts.append(_w2(n))
    return " ".join(parts) + " Rupees Only"


def _fmt_inr(n: float) -> str:
    """Format float in Indian comma style: 2,64,883.20"""
    neg = n < 0
    n = abs(n)
    int_part = int(round(n * 100)) // 100
    frac = int(round(n * 100)) % 100
    s = str(int_part)
    if len(s) <= 3:
        fmt = s
    else:
        fmt = s[-3:]
        s = s[:-3]
        while s:
            chunk = s[-2:] if len(s) >= 2 else s
            fmt = chunk + "," + fmt
            s = s[:-len(chunk)]
    result = f"{fmt}.{frac:02d}"
    return f"-{result}" if neg else result


# ── Billing computation ────────────────────────────────────────────────────────

def _parse_rows(raw) -> list[dict]:
    if not raw:
        return []
    try:
        data = json.loads(raw) if isinstance(raw, str) else raw
        if isinstance(data, dict):
            shift_type = data.get("shift_type")
            if shift_type == "double":
                return (data.get("shift1") or []) + (data.get("shift2") or [])
            if shift_type == "single":
                return data.get("rows") or []
        if isinstance(data, list):
            return data
    except Exception:
        pass
    return []


def _read_billing_snapshot(wl: dict) -> dict | None:
    """Return the frozen billing_snapshot embedded at save time, or None for old records."""
    raw = wl.get("schedule_data")
    if not raw:
        return None
    try:
        data = json.loads(raw) if isinstance(raw, str) else raw
        if isinstance(data, dict):
            return data.get("billing_snapshot")
    except Exception:
        pass
    return None


def _compute_billing(wl: dict, mc: dict) -> dict:
    rental    = float(mc.get("rental_per_month") or 0)
    ot_rate   = float(wl.get("ot_rate") or 0)
    deduction = float(wl.get("deduction") or 0)

    add_op_qty   = float(wl.get("add_operator_qty") or 0)
    add_op_rate  = float(wl.get("add_operator_rate") or 0)
    add_op_amt   = round(add_op_qty * add_op_rate, 2)
    add_acc_qty  = float(wl.get("add_accommodation_qty") or 0)
    add_acc_rate = float(wl.get("add_accommodation_rate") or 0)
    add_acc_amt  = round(add_acc_qty * add_acc_rate, 2)

    snap = _read_billing_snapshot(wl)
    if snap is not None:
        # Use the values that were frozen at save time — immune to later
        # changes in Work Order config (no_of_days, machine_shift_hour).
        qty    = float(snap["qty"])
        ot_hrs = float(snap.get("ot_hours") or 0)
    else:
        # Legacy path: recalculate for records saved before snapshot support.
        rows      = _parse_rows(wl.get("schedule_data"))
        shift_hr  = float(mc.get("machine_shift_hour") or 8)
        no_days   = float(mc.get("no_of_days") or 26)
        work_hrs  = no_days * shift_hr
        actual_hrs = sum(float(r.get("net_time") or 0) for r in rows)
        qty    = round(actual_hrs / work_hrs, 3) if work_hrs > 0 else 0.0
        ot_hrs = round(sum(float(r.get("ot") or 0) for r in rows), 3)

    hiring = round(rental * qty, 2)
    ot_amt = round(ot_hrs * ot_rate, 2)
    return {
        "qty": qty, "rate": rental, "hiring": hiring,
        "ot_hrs": ot_hrs, "ot_rate": ot_rate, "ot_amt": ot_amt,
        "deduction": deduction,
        "add_op_qty": add_op_qty, "add_op_rate": add_op_rate, "add_op_amt": add_op_amt,
        "add_acc_qty": add_acc_qty, "add_acc_rate": add_acc_rate, "add_acc_amt": add_acc_amt,
        "net": max(0.0, hiring + ot_amt - deduction + add_op_amt + add_acc_amt),
    }


def _period_str(billing_month: str, mc: dict | None = None) -> str:
    try:
        dt = datetime.strptime(billing_month.strip(), "%B %Y")
        billing_cycle = (mc or {}).get("billing_cycle") or "Calendar Month"

        if billing_cycle == "Custom":
            _raw = (mc or {}).get("billing_cycle_start_date")
            start_day = 1
            if _raw:
                try:
                    start_day = int(str(_raw).split("-")[2])
                except Exception:
                    start_day = 1
            if start_day > 1:
                from datetime import date as _date, timedelta as _td
                cycle_start = _date(dt.year, dt.month, start_day)
                ny, nm = (dt.year + 1, 1) if dt.month == 12 else (dt.year, dt.month + 1)
                max_next = calendar.monthrange(ny, nm)[1]
                next_cycle_start = _date(ny, nm, min(start_day, max_next))
                cycle_end = next_cycle_start - _td(days=1)
                return (
                    f"{cycle_start.day} {cycle_start.strftime('%b %Y')} to "
                    f"{cycle_end.day} {cycle_end.strftime('%b %Y')}"
                )

        # Calendar Month (default)
        _, last = calendar.monthrange(dt.year, dt.month)
        return f"1 {dt.strftime('%b %Y')} to {last} {dt.strftime('%b %Y')}"
    except Exception:
        return billing_month


def _build_worklog_schedules_html(selected_items: list[dict]) -> str:
    """Build print-ready HTML of work log schedules to append after the invoice."""
    _WL_CSS = """
    <style>
    .wl-page-break { page-break-before: always; }
    .wl-section { padding: 16px 0; }
    .wl-machine-hdr {
        font-size: 11pt; font-weight: 800; color: #1e3a5f;
        border-bottom: 2px solid #1e3a5f; padding-bottom: 5px; margin-bottom: 10px;
    }
    .wl-period { font-size: 8pt; color: #555; margin-bottom: 8px; }
    table.wl-tbl { width: 100%; border-collapse: collapse; font-size: 7.5pt; }
    table.wl-tbl th {
        border: 1px solid #aaa; padding: 4px 6px; background: #1e3a5f;
        color: #fff; font-weight: 700; text-align: center; white-space: nowrap;
    }
    table.wl-tbl td { border: 1px solid #ccc; padding: 3px 6px; text-align: center; }
    table.wl-tbl tr.sun td { background: #dbeafe; color: #1e40af; font-weight: 600; }
    table.wl-tbl tr.ot td.ot-cell { background: #fef3c7; color: #92400e; font-weight: 700; }
    .wl-totals { margin-top: 6px; font-size: 8pt; font-weight: 700;
                 background: #1c1c2e; color: #fff; padding: 6px 10px;
                 border-radius: 0 0 4px 4px; display: flex; gap: 24px; }
    .wl-totals span b { color: #E87722; }
    </style>"""

    sections: list[str] = []
    seen: set[str] = set()

    for item in selected_items:
        if item["type"] != "worklog":
            continue
        mc  = item["mc"]
        wl  = item["wl"]
        mid = mc.get("machine_id", str(item["sl"]))
        if mid in seen:
            continue
        seen.add(mid)

        mlbl   = mc.get("machine_label", "")
        period = item["period"]
        rows   = _parse_rows(wl.get("schedule_data", ""))
        if not rows:
            continue

        # Totals
        total_net = sum(float(r.get("net_time") or 0) for r in rows)
        total_ot  = sum(float(r.get("ot") or 0) for r in rows)
        total_bd  = sum(float(r.get("breakdown_hours") or 0) for r in rows)

        thead = (
            "<tr>"
            "<th>Date</th><th>Day</th>"
            "<th>Start</th><th>End</th>"
            "<th>Net (h)</th><th>OT (h)</th><th>B/D (h)</th>"
            "<th>Start HMR</th><th>End HMR</th><th>Net HMR</th>"
            "<th>Operator</th><th>Remarks</th>"
            "</tr>"
        )

        tbody = ""
        for r in rows:
            wd     = r.get("weekday", "")
            is_sun = wd == "Sunday"
            ot_val = float(r.get("ot") or 0)
            tr_cls = "sun" if is_sun else ""

            def _d(v, fmt=".1f") -> str:
                try:
                    fv = float(v)
                    return f"{fv:{fmt}}" if fv else "—"
                except (TypeError, ValueError):
                    return str(v) if v else "—"

            ot_cell = f"<td class='ot-cell'>{_d(r.get('ot'))}</td>" if ot_val > 0 else f"<td>{_d(r.get('ot'))}</td>"
            tbody += (
                f"<tr class='{tr_cls}'>"
                f"<td>{r.get('date', '—')}</td>"
                f"<td>{wd}</td>"
                f"<td>{r.get('start_time') or '—'}</td>"
                f"<td>{r.get('end_time') or '—'}</td>"
                f"<td>{_d(r.get('net_time'))}</td>"
                f"{ot_cell}"
                f"<td>{_d(r.get('breakdown_hours'))}</td>"
                f"<td>{_d(r.get('start_hmr'))}</td>"
                f"<td>{_d(r.get('end_hmr'))}</td>"
                f"<td>{_d(r.get('net_hmr'))}</td>"
                f"<td style='text-align:left;'>{r.get('operator') or '—'}</td>"
                f"<td style='text-align:left;'>{r.get('remarks') or ''}</td>"
                "</tr>"
            )

        section = (
            f"<div class='wl-page-break wl-section'>"
            f"<div class='wl-machine-hdr'>Work Log — {mlbl}</div>"
            f"<div class='wl-period'>Period: {period}</div>"
            f"<table class='wl-tbl'><thead>{thead}</thead><tbody>{tbody}</tbody></table>"
            f"<div class='wl-totals'>"
            f"<span>Net Time: <b>{total_net:.1f} hrs</b></span>"
            f"<span>OT Hours: <b>{total_ot:.1f} hrs</b></span>"
            f"<span>Breakdown: <b>{total_bd:.1f} hrs</b></span>"
            f"</div></div>"
        )
        sections.append(section)

    if not sections:
        return ""
    return _WL_CSS + "\n".join(sections)


# ── HTML invoice template ──────────────────────────────────────────────────────

_CSS = """
<style>
@page{size:A4;margin:12mm 12mm;}
@media print{
  .no-print{display:none!important;}
  body{background:none!important;padding:0!important;}
  .wrapper{box-shadow:none!important;}
  .wrapper.no-border{border:none!important;}
}
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:Arial,Helvetica,sans-serif;font-size:8.5pt;color:#111;
     background:#ddd;padding:16px;}
.wrapper{width:186mm;margin:0 auto;background:#fff;border:1.5px solid #333;}
.wrapper.no-border{border:none;}
/* header */
.hdr{display:flex;align-items:stretch;border-bottom:1.5px solid #333;}
.hdr-logo{display:flex;align-items:center;justify-content:center;padding:6px 8px 6px 6px;}
.hdr-center{flex:1;padding:6px 10px;display:flex;flex-direction:column;
            align-items:center;justify-content:center;text-align:center;}
.co-name{font-size:18pt;font-weight:900;color:#1A6B1A;letter-spacing:.5px;line-height:1.1;}
.co-sub1{font-size:8.5pt;color:#1A6B1A;font-weight:700;margin-top:3px;}
.co-sub2{font-size:8pt;color:#333;margin-top:1px;}
.co-divider{width:90%;border:none;border-top:1px solid #999;margin:5px 0 4px;}
.co-addr{font-size:7pt;color:#333;font-weight:600;text-align:center;}
.hdr-right{width:130px;display:flex;align-items:center;justify-content:center;
           border-left:1.5px solid #333;padding:8px;}
.ti-box{font-size:11pt;font-weight:900;letter-spacing:2px;border:1.5px solid #333;
        padding:4px 10px;text-align:center;}
.ti-centre-wrap{display:flex;justify-content:center;padding:8px 0;border-bottom:1px solid #888;}
.ti-centre{font-size:13pt;font-weight:900;letter-spacing:3px;border:1.5px solid #333;padding:5px 28px;}
/* two-col sections */
.two-col{display:flex;border-bottom:1px solid #888;}
.two-col .left{flex:1;padding:6px 10px;border-right:1px solid #888;font-size:8pt;}
.two-col .right{flex:1;padding:6px 10px;font-size:8pt;}
.sec-lbl{font-weight:700;text-decoration:underline;margin-bottom:3px;font-size:8pt;}
.fr{display:flex;gap:4px;margin:1px 0;}
.fk{font-weight:700;min-width:100px;flex-shrink:0;}
/* table */
table.inv{width:100%;border-collapse:collapse;font-size:8pt;}
table.inv th{border:1px solid #aaa;padding:4px 5px;background:#eef2f7;
             text-align:center;font-size:7.5pt;font-weight:700;line-height:1.3;}
table.inv td{border:1px solid #aaa;padding:3px 5px;vertical-align:top;}
tr.eq-hdr td{background:#e4ecf5;font-weight:700;font-size:8.5pt;}
tr.subtotal td{font-weight:700;background:#f5f5f5;}
tr.grand td{font-weight:900;font-size:9.5pt;background:#dfe8f4;}
/* footer */
.words{border-top:1px solid #888;padding:5px 10px;font-size:8pt;}
.foot{display:flex;border-top:1px solid #888;}
.foot-left{flex:1;padding:8px 10px;border-right:1px solid #888;font-size:8pt;}
.foot-right{width:160px;padding:8px 10px;font-size:8pt;text-align:right;}
/* print btn */
.pbtn{position:fixed;top:20px;right:20px;background:#2563EB;color:#fff;
      border:none;padding:9px 22px;border-radius:8px;font-size:13px;
      font-weight:700;cursor:pointer;box-shadow:0 4px 12px rgba(37,99,235,.35);}
.pbtn:hover{background:#1D4ED8;}
</style>
"""


def _build_html(
    *,
    inv_no: str,
    inv_date: date,
    wo: dict,
    customer: dict,
    site: dict,
    groups: list[dict],
    tax_type: str,
    tax_on: bool,
    hsn_on: bool,
    hsn_code: str,
    item_code_on: bool,
    notes: str,
    blank_header: bool = False,
) -> str:

    # ── address blocks ─────────────────────────────────────────────────────────
    # Site-level fields (new) take priority; customer fields used as fallback
    cname      = customer.get("customer_name", "")
    bill_addr  = (
        site.get("bill_to_address")
        or customer.get("billing_address")
        or site.get("address")
        or ""
    )
    bill_city  = ", ".join(filter(None, [
        customer.get("city") or site.get("city"),
        customer.get("state") or site.get("state"),
        customer.get("pincode") or site.get("pincode"),
    ]))
    bill_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    bill_state = customer.get("state") or site.get("state") or "—"

    ship_addr  = (
        site.get("ship_to_address")
        or site.get("address")
        or ""
    )
    ship_city  = ", ".join(filter(None, [
        site.get("city"), site.get("state"), site.get("pincode"),
    ]))
    ship_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    ship_state = site.get("state") or "—"

    wo_num    = wo.get("wo_number", "—")
    client_wo = wo.get("client_work_ordernumber") or wo_num
    wo_date   = wo.get("start_date", "—")

    # ── totals ─────────────────────────────────────────────────────────────────
    subtotal = sum(sum(it["amount"] for it in g["items"]) for g in groups)
    tax_rate = 0.18 if tax_on else 0.0
    tax_tot  = round(subtotal * tax_rate, 2)
    grand_ex = subtotal + tax_tot
    grand    = round(grand_ex)
    rnd_off  = round(grand - grand_ex, 2)

    cgst = sgst = igst = 0.0
    if tax_on:
        if tax_type == "CGST/SGST":
            cgst = sgst = round(tax_tot / 2, 2)
        else:
            igst = tax_tot

    # ── column count for colspan ───────────────────────────────────────────────
    n_cols   = 7 + (1 if item_code_on else 0) + (1 if hsn_on else 0)
    span_pre = n_cols - 1   # colspan for all cols before Amount

    # ── line items HTML ────────────────────────────────────────────────────────
    rows_html = ""
    for grp in groups:
        label = grp["machine_label"]
        if grp.get("make") or grp.get("model"):
            label += f" — {' '.join(filter(None,[grp.get('make',''),grp.get('model','')]))}"
        if grp.get("serial"):
            label += f" S/N - {grp['serial']}"

        rows_html += (
            f"<tr class='eq-hdr'><td colspan='{n_cols}' "
            f"style='padding:4px 6px;'>{label}</td></tr>"
        )

        sl = grp["sl_no"]
        for idx, it in enumerate(grp["items"]):
            sl_cell  = str(sl) if idx == 0 else ""
            ic_cell  = (grp.get("item_code") or "") if idx == 0 and item_code_on else ""
            hsn_cell = hsn_code if hsn_on else ""
            tax_cell = "18%" if tax_on else ""
            desc_h   = it["desc"].replace("\n", "<br>")
            amt_fmt  = _fmt_inr(it["amount"])

            ic_td  = f"<td style='border:1px solid #aaa;padding:3px 5px;text-align:center;'>{ic_cell}</td>" if item_code_on else ""
            hsn_td = f"<td style='border:1px solid #aaa;padding:3px 5px;text-align:center;'>{hsn_cell}</td>" if hsn_on else ""

            rows_html += f"""<tr>
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:center;vertical-align:top;'>{sl_cell}</td>
  {ic_td}
  <td style='border:1px solid #aaa;padding:3px 6px;'>{desc_h}</td>
  {hsn_td}
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:center;'>{tax_cell}</td>
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:center;'>{it['uom']}</td>
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:right;'>{it['qty']}</td>
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:right;'>{_fmt_inr(it['rate'])}</td>
  <td style='border:1px solid #aaa;padding:3px 5px;text-align:right;'>{amt_fmt}</td>
</tr>"""

    # ── tax + totals rows ──────────────────────────────────────────────────────
    tax_rows = ""
    if tax_on:
        if tax_type == "CGST/SGST":
            tax_rows = (
                f"<tr><td colspan='{span_pre}' style='border:1px solid #aaa;'></td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;'>CGST 9%</td></tr>"
                .replace("</td></tr>", f"<td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'>{_fmt_inr(cgst)}</td></tr>")
            )
            # Simpler direct approach:
            tax_rows = (
                f"<tr><td colspan='{span_pre}' style='border:1px solid #aaa;padding:3px 6px;'></td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'></td></tr>"
            )
            tax_rows = (
                f"<tr><td colspan='{span_pre-1}' style='border:1px solid #aaa;'></td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;font-size:8pt;'>CGST 9%</td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'>{_fmt_inr(cgst)}</td></tr>"
                f"<tr><td colspan='{span_pre-1}' style='border:1px solid #aaa;'></td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;font-size:8pt;'>SGST 9%</td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'>{_fmt_inr(sgst)}</td></tr>"
            )
        else:
            tax_rows = (
                f"<tr><td colspan='{span_pre-1}' style='border:1px solid #aaa;'></td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;font-size:8pt;'>IGST @ 18%</td>"
                f"<td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'>{_fmt_inr(igst)}</td></tr>"
            )

    rnd_disp = _fmt_inr(rnd_off) if rnd_off != 0 else "—"

    # ── column headers ─────────────────────────────────────────────────────────
    ic_th  = "<th style='width:7%;'>Item<br>code</th>" if item_code_on else ""
    hsn_th = "<th style='width:7%;'>HSN/S<br>AC</th>" if hsn_on else ""

    # ── notes ─────────────────────────────────────────────────────────────────
    notes_html = (
        f"<div style='border-top:1px solid #888;padding:5px 10px;font-size:7.5pt;color:#555;'>"
        f"<b>Notes:</b> {notes}</div>"
    ) if notes else ""

    inv_date_str = inv_date.strftime("%d-%B-%Y") if isinstance(inv_date, date) else str(inv_date)

    if blank_header:
        # Plain spacer — no border, no flex container — leaves room for pre-printed letterhead.
        # Height is calibrated to match the letterhead's bottom edge (~48 mm of printable content).
        _hdr_block   = "<div style='height:48mm;'></div>"
        _ti_block    = '<div class="ti-centre-wrap"><div class="ti-centre">TAX INVOICE</div></div>'
        _wrapper_cls = "wrapper no-border"
    else:
        _hdr_inner = (
            "  <div class='hdr-logo'>\n"
            "    <svg width='90' height='56' viewBox='0 0 90 56' xmlns='http://www.w3.org/2000/svg'>\n"
            "      <rect x='0' y='1' width='40' height='54' fill='#1A6B1A'/>\n"
            "      <text x='20' y='35' fill='white' font-size='18' font-weight='900'\n"
            "            font-family='Arial,Helvetica,sans-serif' text-anchor='middle'>cto</text>\n"
            "      <rect x='45' y='1'  width='43' height='9' fill='#1A6B1A'/>\n"
            "      <rect x='45' y='13' width='43' height='9' fill='#1A6B1A'/>\n"
            "      <rect x='45' y='25' width='43' height='9' fill='#1A6B1A'/>\n"
            "      <rect x='45' y='37' width='43' height='9' fill='#1A6B1A'/>\n"
            "      <rect x='45' y='47' width='43' height='8' fill='#1A6B1A'/>\n"
            "    </svg>\n"
            "  </div>\n"
            "  <div class='hdr-center'>\n"
            "    <div class='co-name'>CTO LOGISTICS &amp; INFRA</div>\n"
            "    <div class='co-sub1'>(CTO GROUP)</div>\n"
            "    <div class='co-sub2'>(LOGISTICS &amp; INFRA EQUIPMENTS)</div>\n"
            "    <hr class='co-divider'/>\n"
            "    <div class='co-addr'>\n"
            f"      B-202, STEEL CHAMBERS, STEEL MARKET ROAD, PLOT NO. 514, KALAMBOLI - 410 208, DIST. RAIGAD\n"
            f"      &nbsp; Tel.: {_CO['tel']} &nbsp; E-mail: {_CO['email']}\n"
            "    </div>\n"
            "  </div>"
        )
        _hdr_block   = f"<div class='hdr'>\n{_hdr_inner}\n  <div class='hdr-right'><div class='ti-box'>TAX<br>INVOICE</div></div>\n</div>"
        _ti_block    = ""
        _wrapper_cls = "wrapper"

    _print_hint = (
        "<p class='no-print' style='text-align:center;font-size:11px;color:#6b7280;"
        "margin:6px 0 10px;'>"
        "⚠ In the print dialog, uncheck <b>Headers and Footers</b> for a clean output."
        "</p>"
    ) if blank_header else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>Tax Invoice — {inv_no}</title>
{_CSS}
</head>
<body>
<button class="pbtn no-print" onclick="window.print()">🖨 Print / Save PDF</button>
{_print_hint}
<div class="{_wrapper_cls}">

<!-- ── Company header / letterhead spacer ── -->
{_hdr_block}
{_ti_block}
<!-- ── Company details | Invoice meta ── -->
<div class="two-col">
  <div class="left">
    <b>{_CO['name']}</b><br>
    {_CO['addr1']}<br>{_CO['addr2']}<br>{_CO['addr3']}<br>{_CO['addr4']}<br>
    <div class="fr"><span class="fk">GSTIN/UIN:</span><span>{_CO['gstin']}</span></div>
    <div class="fr"><span class="fk">State Name :</span>
      <span>{_CO['state']}, Code : {_CO['state_code']}</span></div>
    <div class="fr"><span class="fk">E-Mail :</span><span>{_CO['email']}</span></div>
  </div>
  <div class="right">
    <div class="fr"><span class="fk">Invoice No. :-</span><span><b>{inv_no}</b></span></div>
    <div class="fr"><span class="fk">Dated :-</span><span>{inv_date_str}</span></div>
    <br>
    <div class="fr"><span class="fk">Work Order No. -</span><span>{client_wo}</span></div>
    <div class="fr"><span class="fk">Word Order Dt. -</span><span>{wo_date}</span></div>
  </div>
</div>

<!-- ── Ship to | Bill to ── -->
<div class="two-col">
  <div class="left">
    <div class="sec-lbl">Consignee (Ship to)</div>
    <b>{cname}</b><br>
    {ship_addr}<br>{ship_city}<br>
    <div class="fr"><span class="fk">GSTIN/UIN :</span><span>{ship_gst}</span></div>
    <div class="fr"><span class="fk">State Name :</span><span>{ship_state}</span></div>
  </div>
  <div class="right">
    <div class="sec-lbl">Buyer (Bill to)</div>
    <b>{cname}</b><br>
    {bill_addr}<br>{bill_city}<br>
    <div class="fr"><span class="fk">GSTIN/UIN :</span><span>{bill_gst}</span></div>
    <div class="fr"><span class="fk">State Name :</span><span>{bill_state}</span></div>
  </div>
</div>

<!-- ── Line items ── -->
<table class="inv">
<thead>
  <tr>
    <th style='width:4%;'>Sl<br>No.</th>
    {ic_th}
    <th>Description of Services</th>
    {hsn_th}
    <th style='width:5%;'>Tax<br>rate</th>
    <th style='width:6%;'>UOM</th>
    <th style='width:8%;'>Quantity</th>
    <th style='width:11%;'>Rate</th>
    <th style='width:12%;'>Amount (INR)</th>
  </tr>
</thead>
<tbody>
{rows_html}
<!-- subtotal -->
<tr class="subtotal">
  <td colspan='{span_pre}' style='border:1px solid #aaa;padding:4px 6px;'>Total taxable vale</td>
  <td style='border:1px solid #aaa;padding:4px 6px;text-align:right;'>{_fmt_inr(subtotal)}</td>
</tr>
{tax_rows}
<!-- round off -->
<tr>
  <td colspan='{span_pre}' style='border:1px solid #aaa;padding:3px 6px;'>Round Off</td>
  <td style='border:1px solid #aaa;padding:3px 6px;text-align:right;'>{rnd_disp}</td>
</tr>
<!-- grand total -->
<tr class="grand">
  <td colspan='{span_pre}' style='border:1px solid #aaa;padding:5px 6px;'>Grand total</td>
  <td style='border:1px solid #aaa;padding:5px 6px;text-align:right;'>{_fmt_inr(grand)}</td>
</tr>
</tbody>
</table>

<!-- ── Amount in words ── -->
<div class="words">
  <b>Amount Chargeable (in words)</b> - INR {_num_to_words_inr(grand)}
</div>

{notes_html}

<!-- ── Bank details + signature ── -->
<div class="foot">
  <div class="foot-left">
    <b>Company Bank Details</b><br><br>
    <div class="fr"><span class="fk">A/c Holder Name :</span><span>{_BANK['holder']}</span></div>
    <div class="fr"><span class="fk">Bank Name :</span><span>{_BANK['bank']}</span></div>
    <div class="fr"><span class="fk">A/c No. :</span><span>{_BANK['account']}</span></div>
    <div class="fr"><span class="fk">Branch &amp; IFSC Code :</span><span>{_BANK['ifsc']}</span></div>
  </div>
  <div class="foot-right">
    For {_CO['name']}<br><br><br><br><br>
    Authorised Signatory
  </div>
</div>

</div><!-- /wrapper -->
</body></html>"""


# ── Word (.docx) invoice builder ───────────────────────────────────────────────

def _build_docx(
    *,
    inv_no: str,
    inv_date: date,
    wo: dict,
    customer: dict,
    site: dict,
    groups: list[dict],
    tax_type: str,
    tax_on: bool,
    hsn_on: bool,
    hsn_code: str,
    item_code_on: bool,
    notes: str,
    blank_header: bool = False,
) -> bytes:
    """Return a .docx invoice as raw bytes. Mirrors _build_html structure."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── helpers ────────────────────────────────────────────────────────────────

    def _shd(cell, hex_fill: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for x in tcPr.findall(qn("w:shd")):
            tcPr.remove(x)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_fill.lstrip("#").upper())
        tcPr.append(shd)

    def _cwrite(cell, text: str, size: float = 8.5, bold: bool = False,
                align: str = "left", first: bool = True) -> None:
        """Write possibly multi-line text into a cell paragraph."""
        lines = str(text if text is not None else "").split("\n")
        for li, line in enumerate(lines):
            use_first = first and li == 0
            para = cell.paragraphs[0] if use_first else cell.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            if align == "center":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(line)
            run.font.size = Pt(size)
            run.bold = bold

    def _lv(cell, label: str, value: str) -> None:
        """Add a bold-label : value paragraph to a cell."""
        para = cell.add_paragraph()
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)
        r1 = para.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(8.5)
        r2 = para.add_run(value or "—")
        r2.font.size = Pt(8.5)

    def _merge_row(row, from_col: int, to_col: int) -> None:
        """Merge cells from_col..to_col (inclusive) in row."""
        for j in range(from_col + 1, to_col + 1):
            row.cells[from_col].merge(row.cells[j])

    def _mktbl(rows: int, cols: int, widths_mm: list) -> object:
        t = doc.add_table(rows=rows, cols=cols)
        t.style   = "Table Grid"
        t.autofit = False
        for i, col in enumerate(t.columns):
            w = widths_mm[i] if i < len(widths_mm) else widths_mm[-1]
            for cell in col.cells:
                cell.width = Mm(w)
        return t

    # ── pre-compute totals ─────────────────────────────────────────────────────
    subtotal  = sum(sum(it["amount"] for it in g["items"]) for g in groups)
    tax_tot   = round(subtotal * 0.18, 2) if tax_on else 0.0
    grand_ex  = subtotal + tax_tot
    grand     = round(grand_ex)
    rnd_off   = round(grand - grand_ex, 2)
    cgst = sgst = igst = 0.0
    if tax_on:
        if tax_type == "CGST/SGST":
            cgst = sgst = round(tax_tot / 2, 2)
        else:
            igst = tax_tot

    inv_date_str = inv_date.strftime("%d-%B-%Y") if isinstance(inv_date, date) else str(inv_date)
    wo_num    = wo.get("wo_number", "—")
    client_wo = wo.get("client_work_ordernumber") or wo_num
    wo_date   = str(wo.get("start_date", "—"))
    cname     = customer.get("customer_name", "")

    bill_addr  = (site.get("bill_to_address") or customer.get("billing_address")
                  or site.get("address") or "")
    bill_city  = ", ".join(filter(None, [
        customer.get("city") or site.get("city"),
        customer.get("state") or site.get("state"),
        customer.get("pincode") or site.get("pincode"),
    ]))
    bill_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    bill_state = customer.get("state") or site.get("state") or "—"

    ship_addr  = site.get("ship_to_address") or site.get("address") or ""
    ship_city  = ", ".join(filter(None, [
        site.get("city"), site.get("state"), site.get("pincode"),
    ]))
    ship_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    ship_state = site.get("state") or "—"

    # ── document setup ─────────────────────────────────────────────────────────
    doc = Document()
    ns  = doc.styles["Normal"]
    ns.font.name  = "Arial"
    ns.font.size  = Pt(9)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(2)

    sec = doc.sections[0]
    sec.page_width    = Mm(210)
    sec.page_height   = Mm(297)
    sec.left_margin   = Mm(12)
    sec.right_margin  = Mm(12)
    sec.top_margin    = Mm(12)
    sec.bottom_margin = Mm(12)

    PAGE_W = 186  # usable page width in mm

    # ── 1. Company header ──────────────────────────────────────────────────────
    t_hdr = _mktbl(1, 2, [148, 38])

    lc = t_hdr.cell(0, 0)
    if blank_header:
        # Leave left cell blank — space for pre-printed letterhead (~42mm from top)
        p_blank = lc.paragraphs[0]
        p_blank.paragraph_format.space_before = Pt(85)
        p_blank.paragraph_format.space_after  = Pt(0)
    else:
        p  = lc.paragraphs[0]
        r  = p.add_run("CTO LOGISTICS & INFRA")
        r.bold = True
        r.font.size  = Pt(16)
        r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x1A)

        p2 = lc.add_paragraph("(CTO GROUP)")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.runs[0]
        r2.bold = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x1A, 0x6B, 0x1A)

        _cwrite(lc, "(LOGISTICS & INFRA EQUIPMENTS)", size=7.5, first=False)
        addr_ln = (
            f"B-202, STEEL CHAMBERS, STEEL MARKET ROAD, PLOT NO. 514, KALAMBOLI - 410 208, "
            f"DIST. RAIGAD  ·  Tel.: {_CO['tel']}  ·  {_CO['email']}"
        )
        p_addr = lc.add_paragraph(addr_ln)
        p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_addr.runs[0].font.size = Pt(7.5)
        p_addr.paragraph_format.space_before = Pt(2)

    rc = t_hdr.cell(0, 1)
    if not blank_header:
        rc.vertical_alignment = 1  # CENTER
        p_ti = rc.paragraphs[0]
        p_ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ti = p_ti.add_run("TAX\nINVOICE")
        r_ti.bold = True
        r_ti.font.size = Pt(13)

    if blank_header:
        p_ti_c = doc.add_paragraph()
        p_ti_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ti_c.paragraph_format.space_before = Pt(6)
        p_ti_c.paragraph_format.space_after  = Pt(6)
        r_ti_c = p_ti_c.add_run("TAX INVOICE")
        r_ti_c.bold = True
        r_ti_c.font.size = Pt(16)

    # ── 2. GSTIN / invoice meta ────────────────────────────────────────────────
    t_meta = _mktbl(1, 2, [93, 93])

    lm = t_meta.cell(0, 0)
    _cwrite(lm, _CO["name"], bold=True)
    for line in [_CO["addr1"], _CO["addr2"], _CO["addr3"], _CO["addr4"]]:
        _cwrite(lm, line, size=8.5, first=False)
    _lv(lm, "GSTIN/UIN",   _CO["gstin"])
    _lv(lm, "State Name",  f"{_CO['state']}, Code: {_CO['state_code']}")
    _lv(lm, "E-Mail",      _CO["email"])

    rm = t_meta.cell(0, 1)
    _cwrite(rm, "")  # empty anchor paragraph
    _lv(rm, "Invoice No.", inv_no)
    _lv(rm, "Dated",       inv_date_str)
    rm.add_paragraph()
    _lv(rm, "Work Order No.", client_wo)
    _lv(rm, "Work Order Dt.", wo_date)

    # ── 3. Ship to / Bill to ──────────────────────────────────────────────────
    t_addr = _mktbl(1, 2, [93, 93])

    sc = t_addr.cell(0, 0)
    p  = sc.paragraphs[0]
    r  = p.add_run("Consignee (Ship to)")
    r.bold = True; r.underline = True; r.font.size = Pt(8.5)
    _cwrite(sc, cname, bold=True, size=8.5, first=False)
    for ln in filter(None, [ship_addr, ship_city]):
        _cwrite(sc, ln, size=8.5, first=False)
    _lv(sc, "GSTIN/UIN",  ship_gst)
    _lv(sc, "State Name", ship_state)

    bc = t_addr.cell(0, 1)
    p  = bc.paragraphs[0]
    r  = p.add_run("Buyer (Bill to)")
    r.bold = True; r.underline = True; r.font.size = Pt(8.5)
    _cwrite(bc, cname, bold=True, size=8.5, first=False)
    for ln in filter(None, [bill_addr, bill_city]):
        _cwrite(bc, ln, size=8.5, first=False)
    _lv(bc, "GSTIN/UIN",  bill_gst)
    _lv(bc, "State Name", bill_state)

    # ── 4. Line items table ────────────────────────────────────────────────────
    col_hdrs: list[str] = ["Sl\nNo."]
    col_w:    list[float] = [8.0]
    if item_code_on:
        col_hdrs.append("Item\nCode"); col_w.append(14.0)
    col_hdrs.append("Description of Services"); col_w.append(0.0)  # fill later
    if hsn_on:
        col_hdrs.append("HSN/\nSAC"); col_w.append(14.0)
    col_hdrs += ["Tax\nRate", "UOM", "Quantity",  "Rate",  "Amount\n(INR)"]
    col_w    += [9.0,         11.0,   16.0,        21.0,    22.0]

    desc_i = col_hdrs.index("Description of Services")
    col_w[desc_i] = PAGE_W - sum(w for w in col_w if w > 0)

    n_cols = len(col_hdrs)
    t_items = _mktbl(1, n_cols, col_w)

    # header row
    for i, hdr in enumerate(col_hdrs):
        cell = t_items.rows[0].cells[i]
        _shd(cell, "EEF2F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.bold = True; r.font.size = Pt(8.0)

    # line item rows
    for grp in groups:
        label = grp["machine_label"]
        if grp.get("make") or grp.get("model"):
            label += " — " + " ".join(filter(None, [grp.get("make",""), grp.get("model","")]))
        if grp.get("serial"):
            label += f" S/N - {grp['serial']}"

        eq_row = t_items.add_row()
        _merge_row(eq_row, 0, n_cols - 1)
        _shd(eq_row.cells[0], "E4ECF5")
        r = eq_row.cells[0].paragraphs[0].add_run(label)
        r.bold = True; r.font.size = Pt(9.0)

        sl = grp["sl_no"]
        ic = grp.get("item_code", "") or ""

        for idx, it in enumerate(grp["items"]):
            irow = t_items.add_row()
            ci   = 0
            _cwrite(irow.cells[ci], str(sl) if idx == 0 else "", align="center"); ci += 1
            if item_code_on:
                _cwrite(irow.cells[ci], ic if idx == 0 else "", align="center"); ci += 1
            _cwrite(irow.cells[ci], it["desc"], align="left"); ci += 1
            if hsn_on:
                _cwrite(irow.cells[ci], hsn_code, align="center"); ci += 1
            _cwrite(irow.cells[ci], "18%" if tax_on else "", align="center"); ci += 1
            _cwrite(irow.cells[ci], it["uom"],           align="center"); ci += 1
            _cwrite(irow.cells[ci], str(it["qty"]),       align="right");  ci += 1
            _cwrite(irow.cells[ci], _fmt_inr(it["rate"]), align="right");  ci += 1
            _cwrite(irow.cells[ci], _fmt_inr(it["amount"]), align="right")

    # subtotal
    sr = t_items.add_row()
    _merge_row(sr, 0, n_cols - 2)
    _shd(sr.cells[0], "F5F5F5"); _shd(sr.cells[-1], "F5F5F5")
    r = sr.cells[0].paragraphs[0].add_run("Total taxable value")
    r.bold = True; r.font.size = Pt(8.5)
    _cwrite(sr.cells[-1], _fmt_inr(subtotal), bold=True, align="right")

    # tax rows
    if tax_on:
        if tax_type == "CGST/SGST":
            for lbl, val in [("CGST 9%", cgst), ("SGST 9%", sgst)]:
                tr_ = t_items.add_row()
                if n_cols > 2:
                    _merge_row(tr_, 0, n_cols - 3)
                tr_.cells[-2].paragraphs[0].add_run(lbl).font.size = Pt(8.5)
                _cwrite(tr_.cells[-1], _fmt_inr(val), align="right")
        else:
            tr_ = t_items.add_row()
            if n_cols > 2:
                _merge_row(tr_, 0, n_cols - 3)
            tr_.cells[-2].paragraphs[0].add_run("IGST @ 18%").font.size = Pt(8.5)
            _cwrite(tr_.cells[-1], _fmt_inr(igst), align="right")

    # round off
    rr = t_items.add_row()
    if n_cols > 2:
        _merge_row(rr, 0, n_cols - 3)
    rr.cells[-2].paragraphs[0].add_run("Round Off").font.size = Pt(8.5)
    _cwrite(rr.cells[-1], _fmt_inr(rnd_off) if rnd_off != 0 else "—", align="right")

    # grand total
    gr = t_items.add_row()
    _merge_row(gr, 0, n_cols - 2)
    _shd(gr.cells[0], "DFE8F4"); _shd(gr.cells[-1], "DFE8F4")
    r = gr.cells[0].paragraphs[0].add_run("Grand Total")
    r.bold = True; r.font.size = Pt(10.0)
    p_ga = gr.cells[-1].paragraphs[0]
    p_ga.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_ga.add_run(_fmt_inr(grand))
    r.bold = True; r.font.size = Pt(10.0)

    # ── 5. Amount in words ─────────────────────────────────────────────────────
    p_w = doc.add_paragraph()
    p_w.paragraph_format.space_before = Pt(4)
    r1 = p_w.add_run("Amount Chargeable (in words): ")
    r1.bold = True; r1.font.size = Pt(8.5)
    r2 = p_w.add_run(f"INR {_num_to_words_inr(grand)}")
    r2.font.size = Pt(8.5)

    # ── 6. Notes ──────────────────────────────────────────────────────────────
    if notes:
        p_n = doc.add_paragraph()
        r1 = p_n.add_run("Notes: ")
        r1.bold = True; r1.font.size = Pt(8.5)
        r2 = p_n.add_run(notes)
        r2.font.size = Pt(8.5)

    # ── 7. Bank details + authorised signatory ────────────────────────────────
    t_foot = _mktbl(1, 2, [126, 60])

    bk = t_foot.cell(0, 0)
    _cwrite(bk, "Company Bank Details", bold=True)
    bk.add_paragraph()
    _lv(bk, "A/c Holder Name",    _BANK["holder"])
    _lv(bk, "Bank Name",          _BANK["bank"])
    _lv(bk, "A/c No.",            _BANK["account"])
    _lv(bk, "Branch & IFSC Code", _BANK["ifsc"])

    sg = t_foot.cell(0, 1)
    p_sg = sg.paragraphs[0]
    p_sg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_sg.add_run(f"For {_CO['name']}")
    r.bold = True; r.font.size = Pt(8.5)
    for _ in range(4):
        p_sp = sg.add_paragraph("")
        p_sp.paragraph_format.space_after = Pt(8)
    p_as = sg.add_paragraph("Authorised Signatory")
    p_as.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p_as.runs:
        p_as.runs[0].font.size = Pt(8.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── PDF (.pdf) invoice builder ────────────────────────────────────────────────

def _build_pdf_bytes(
    *,
    inv_no: str,
    inv_date: date,
    wo: dict,
    customer: dict,
    site: dict,
    groups: list[dict],
    tax_type: str,
    tax_on: bool,
    hsn_on: bool,
    hsn_code: str,
    item_code_on: bool,
    notes: str,
    blank_header: bool = False,
) -> bytes:
    """Return a .pdf invoice as raw bytes using fpdf2 (no system-library deps)."""
    from fpdf import FPDF

    # ── shared data extraction ──────────────────────────────────────────────────
    cname     = customer.get("customer_name", "")
    bill_addr = (
        site.get("bill_to_address") or customer.get("billing_address")
        or site.get("address") or ""
    )
    bill_city = ", ".join(filter(None, [
        customer.get("city") or site.get("city"),
        customer.get("state") or site.get("state"),
        customer.get("pincode") or site.get("pincode"),
    ]))
    bill_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    bill_state = customer.get("state") or site.get("state") or "—"
    ship_addr  = site.get("ship_to_address") or site.get("address") or ""
    ship_city  = ", ".join(filter(None, [
        site.get("city"), site.get("state"), site.get("pincode"),
    ]))
    ship_gst   = site.get("gst_number") or customer.get("gst_number") or "—"
    ship_state = site.get("state") or "—"
    wo_num     = wo.get("wo_number", "—")
    client_wo  = wo.get("client_work_ordernumber") or wo_num
    wo_date    = wo.get("start_date", "—")

    subtotal = sum(sum(it["amount"] for it in g["items"]) for g in groups)
    tax_rate = 0.18 if tax_on else 0.0
    tax_tot  = round(subtotal * tax_rate, 2)
    grand_ex = subtotal + tax_tot
    grand    = round(grand_ex)
    rnd_off  = round(grand - grand_ex, 2)
    cgst = sgst = igst = 0.0
    if tax_on:
        if tax_type == "CGST/SGST":
            cgst = sgst = round(tax_tot / 2, 2)
        else:
            igst = tax_tot

    inv_date_str = (
        inv_date.strftime("%d-%B-%Y") if isinstance(inv_date, date) else str(inv_date)
    )

    # ── fpdf2 helpers ───────────────────────────────────────────────────────────
    def _s(t: object) -> str:
        """Sanitize text to windows-1252; replace anything still unrepresentable."""
        return (
            str(t or "")
            .encode("windows-1252", errors="replace")
            .decode("windows-1252")
        )

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_doc_option("core_fonts_encoding", "windows-1252")
    pdf.set_margins(10, 10, 10)
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.add_page()

    LM: float = 10
    W:  float = 190

    def _hline(y: float | None = None, thick: bool = False) -> None:
        yy = y if y is not None else pdf.get_y()
        r = 51 if thick else 136
        pdf.set_draw_color(r, r, r)
        pdf.line(LM, yy, LM + W, yy)
        pdf.set_draw_color(0, 0, 0)

    # ── HEADER: CTO logo | company name | TAX INVOICE box ───────────────────────
    y0      = pdf.get_y()
    _GRN    = (26, 107, 26)   # #1A6B1A
    logo_w  = 30              # total logo zone width (mm)
    ti_w    = 42
    co_w    = W - logo_w - ti_w   # company name width
    logo_x  = LM
    co_x    = LM + logo_w
    ti_x    = LM + logo_w + co_w
    hdr_h   = 32 if blank_header else 22   # blank: clear pre-printed letterhead (~42mm from page top)

    if not blank_header:
        # ── logo: green filled box with "cto" ────
        box_w, box_h = 13, hdr_h
        pdf.set_fill_color(*_GRN)
        pdf.rect(logo_x, y0, box_w, box_h, "F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_xy(logo_x, y0 + box_h / 2 - 3)
        pdf.cell(box_w, 6, "cto", align="C")

        # ── logo: 5 horizontal green bars ────
        bar_x = logo_x + box_w + 1.5
        bar_w = logo_w - box_w - 2.5
        bar_h = 3.2
        gap   = (hdr_h - 5 * bar_h) / 6
        pdf.set_fill_color(*_GRN)
        for i in range(5):
            by = y0 + gap + i * (bar_h + gap)
            pdf.rect(bar_x, by, bar_w, bar_h, "F")

        # ── company name ────
        pdf.set_xy(co_x, y0 + 1)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(*_GRN)
        pdf.cell(co_w, 6, "CTO LOGISTICS & INFRA", align="C")

        pdf.set_xy(co_x, y0 + 7)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(*_GRN)
        pdf.cell(co_w, 4, "(CTO GROUP)", align="C")

        pdf.set_xy(co_x, y0 + 11)
        pdf.set_font("Helvetica", "", 6.5)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(co_w, 3.5, "(LOGISTICS & INFRA EQUIPMENTS)", align="C")

        pdf.set_draw_color(160, 160, 160)
        pdf.line(co_x + 2, y0 + 15, co_x + co_w - 2, y0 + 15)
        pdf.set_draw_color(0, 0, 0)

        pdf.set_xy(co_x, y0 + 15.5)
        pdf.set_font("Helvetica", "B", 6)
        pdf.set_text_color(51, 51, 51)
        addr_str = (
            "B-202, STEEL CHAMBERS, STEEL MARKET ROAD, PLOT NO. 514,"
            " KALAMBOLI - 410 208, DIST. RAIGAD"
        )
        pdf.multi_cell(co_w, 3, addr_str, align="C", new_x="LEFT", new_y="NEXT")
        pdf.set_xy(co_x, pdf.get_y())
        pdf.set_font("Helvetica", "", 6)
        pdf.cell(co_w, 3, f"Tel.: {_CO['tel']}   E-mail: {_CO['email']}", align="C")

    if not blank_header:
        # ── TAX INVOICE box (top-right, beside logo and company name) ────
        pdf.set_draw_color(51, 51, 51)
        pdf.rect(ti_x, y0, ti_w, hdr_h)
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(ti_x, y0 + 4)
        pdf.cell(ti_w, 6, "TAX", align="C")
        pdf.set_xy(ti_x, y0 + 11)
        pdf.cell(ti_w, 6, "INVOICE", align="C")

    hdr_end = y0 + hdr_h
    pdf.set_y(hdr_end)
    _hline(hdr_end, thick=True)

    if blank_header:
        # ── TAX INVOICE: centred title below the blank letterhead space ────
        _ti_w, _ti_h = 64, 10
        _ti_x = LM + (W - _ti_w) / 2
        pdf.set_draw_color(51, 51, 51)
        pdf.rect(_ti_x, hdr_end + 2, _ti_w, _ti_h)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(0, 0, 0)
        pdf.set_xy(_ti_x, hdr_end + 4)
        pdf.cell(_ti_w, 6, "TAX INVOICE", align="C")
        s1y = hdr_end + _ti_h + 5
    else:
        s1y = hdr_end + 2

    # ── SECTION 1: company details | invoice meta ────────────────────────────────
    half = W / 2

    # left column — company GST details
    pdf.set_xy(LM, s1y)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "B", 8)
    pdf.cell(half - 2, 5, _CO["name"], ln=True)
    pdf.set_x(LM)
    pdf.set_font("Helvetica", "", 7.5)
    for ln_txt in [_CO["addr1"], _CO["addr2"], _CO["addr3"], _CO["addr4"]]:
        pdf.cell(half - 2, 4, ln_txt, ln=True)
        pdf.set_x(LM)
    for lbl, val in [
        ("GSTIN/UIN:", _CO["gstin"]),
        ("State Name:", f"{_CO['state']}, Code: {_CO['state_code']}"),
        ("E-Mail:", _CO["email"]),
    ]:
        pdf.set_font("Helvetica", "B", 7.5)
        pdf.cell(26, 4.5, lbl)
        pdf.set_font("Helvetica", "", 7.5)
        pdf.cell(half - 28, 4.5, val, ln=True)
        pdf.set_x(LM)
    left_y1 = pdf.get_y()

    # right column — invoice meta
    rx = LM + half + 2
    pdf.set_xy(rx, s1y)
    pdf.set_font("Helvetica", "B", 8)
    pdf.cell(half - 4, 5, f"Invoice No.: {inv_no}", ln=True)
    pdf.set_x(rx)
    pdf.set_font("Helvetica", "", 8)
    for txt in [f"Dated: {inv_date_str}", "", f"Work Order No.: {client_wo}",
                f"Work Order Dt.: {wo_date}"]:
        pdf.cell(half - 4, 5, txt, ln=True)
        pdf.set_x(rx)
    right_y1 = pdf.get_y()

    divx = LM + half
    bot1 = max(left_y1, right_y1)
    pdf.set_draw_color(136, 136, 136)
    pdf.line(divx, s1y, divx, bot1)
    pdf.set_y(bot1)
    _hline()

    # ── SECTION 2: ship to | bill to ────────────────────────────────────────────
    s2y = pdf.get_y() + 1.5

    def _addr_block(x: float, heading: str, addr: str, city: str,
                    gst: str, state: str) -> float:
        """Draw one address column; returns bottom Y."""
        pdf.set_xy(x, s2y)
        pdf.set_font("Helvetica", "BU", 8)
        pdf.cell(half - 2, 5, heading, ln=True)
        pdf.set_x(x)
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(half - 2, 4.5, cname, ln=True)
        pdf.set_x(x)
        pdf.set_font("Helvetica", "", 7.5)
        for part in (addr + "\n" + city).split("\n"):
            if part.strip():
                pdf.cell(half - 2, 4, part.strip(), ln=True)
                pdf.set_x(x)
        for lbl, val in [("GSTIN/UIN:", gst), ("State Name:", state)]:
            pdf.set_font("Helvetica", "B", 7.5)
            pdf.cell(22, 4.5, lbl)
            pdf.set_font("Helvetica", "", 7.5)
            pdf.cell(half - 24, 4.5, val, ln=True)
            pdf.set_x(x)
        return pdf.get_y()

    left_y2  = _addr_block(LM,      "Consignee (Ship to)", ship_addr, ship_city, ship_gst, ship_state)
    right_y2 = _addr_block(LM + half + 2, "Buyer (Bill to)", bill_addr, bill_city, bill_gst, bill_state)

    bot2 = max(left_y2, right_y2)
    pdf.set_draw_color(136, 136, 136)
    pdf.line(divx, s2y, divx, bot2)
    pdf.set_y(bot2)
    _hline()

    # ── LINE ITEMS TABLE ─────────────────────────────────────────────────────────
    pdf.ln(1)

    sl_w  = 8
    ic_w  = 15 if item_code_on else 0
    hn_w  = 14 if hsn_on else 0
    tr_w  = 10
    um_w  = 12
    qt_w  = 14
    rt_w  = 22
    am_w  = 24
    dc_w  = W - sl_w - ic_w - hn_w - tr_w - um_w - qt_w - rt_w - am_w
    rh    = 5

    def _th(txt: str, w: float, al: str = "C") -> None:
        pdf.set_fill_color(238, 242, 247)
        pdf.set_font("Helvetica", "B", 7.5)
        pdf.cell(w, rh, _s(txt), border=1, align=al, fill=True)

    def _td(txt: str, w: float, al: str = "L", bold: bool = False,
            rgb: tuple | None = None) -> None:
        if rgb:
            pdf.set_fill_color(*rgb)
        pdf.set_font("Helvetica", "B" if bold else "", 8)
        pdf.cell(w, rh, _s(txt), border=1, align=al, fill=bool(rgb))

    _th("Sl No.", sl_w)
    if item_code_on:
        _th("Item Code", ic_w)
    _th("Description of Services", dc_w, "L")
    if hsn_on:
        _th("HSN/SAC", hn_w)
    _th("Tax Rate", tr_w)
    _th("UOM", um_w)
    _th("Quantity", qt_w)
    _th("Rate", rt_w)
    _th("Amount (INR)", am_w)
    pdf.ln()

    GRP_RGB = (228, 236, 245)
    SUB_RGB = (245, 245, 245)
    GRD_RGB = (223, 232, 244)

    for grp in groups:
        lbl = grp["machine_label"]
        if grp.get("make") or grp.get("model"):
            lbl += " - " + " ".join(filter(None, [grp.get("make",""), grp.get("model","")]))
        if grp.get("serial"):
            lbl += f" S/N {grp['serial']}"
        pdf.set_fill_color(*GRP_RGB)
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.cell(W, rh, lbl, border=1, align="L", fill=True, ln=True)

        for idx, it in enumerate(grp["items"]):
            desc = (it["desc"] or "")[:90]
            _td(str(grp["sl_no"]) if idx == 0 else "", sl_w, "C")
            if item_code_on:
                _td((grp.get("item_code") or "") if idx == 0 else "", ic_w, "C")
            _td(desc, dc_w)
            if hsn_on:
                _td(hsn_code, hn_w, "C")
            _td("18%" if tax_on else "", tr_w, "C")
            _td(it["uom"], um_w, "C")
            _td(str(it["qty"]), qt_w, "R")
            _td(_fmt_inr(it["rate"]), rt_w, "R")
            _td(_fmt_inr(it["amount"]), am_w, "R")
            pdf.ln()

    # totals
    span_w = W - am_w
    pdf.set_fill_color(*SUB_RGB)
    pdf.set_font("Helvetica", "B", 8)
    pdf.cell(span_w, rh, "Total taxable value", border=1, align="L", fill=True)
    pdf.cell(am_w,   rh, _fmt_inr(subtotal), border=1, align="R", fill=True, ln=True)

    if tax_on:
        tax_items = (
            [("CGST 9%", cgst), ("SGST 9%", sgst)]
            if tax_type == "CGST/SGST"
            else [("IGST @ 18%", igst)]
        )
        for t_lbl, t_amt in tax_items:
            pdf.set_font("Helvetica", "", 7.5)
            pdf.cell(span_w - rt_w, rh, "", border=1)
            pdf.cell(rt_w,          rh, t_lbl, border=1, align="L")
            pdf.cell(am_w,          rh, _fmt_inr(t_amt), border=1, align="R", ln=True)

    pdf.set_font("Helvetica", "", 8)
    pdf.cell(span_w, rh, "Round Off", border=1, align="L")
    pdf.cell(am_w,   rh, _fmt_inr(rnd_off) if rnd_off != 0 else "—", border=1, align="R", ln=True)

    pdf.set_fill_color(*GRD_RGB)
    pdf.set_font("Helvetica", "B", 9.5)
    pdf.cell(span_w, rh + 1, "Grand total", border=1, align="L", fill=True)
    pdf.cell(am_w,   rh + 1, _fmt_inr(grand), border=1, align="R", fill=True, ln=True)

    # ── Amount in words ──────────────────────────────────────────────────────────
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(0, 0, 0)
    pdf.set_x(LM)
    pdf.multi_cell(W, 5, f"Amount Chargeable (in words) — INR {_num_to_words_inr(grand)}")

    if notes:
        pdf.set_x(LM)
        pdf.set_text_color(80, 80, 80)
        pdf.multi_cell(W, 5, f"Notes: {notes}")
        pdf.set_text_color(0, 0, 0)

    # ── Footer: bank details | signatory ─────────────────────────────────────────
    pdf.ln(2)
    _hline()
    foot_y = pdf.get_y() + 2
    bank_w = 115

    pdf.set_xy(LM, foot_y)
    pdf.set_font("Helvetica", "B", 8)
    pdf.cell(bank_w, 5, "Company Bank Details", ln=True)
    pdf.set_x(LM)
    pdf.ln(1)
    for f_lbl, f_val in [
        ("A/c Holder Name:", _BANK["holder"]),
        ("Bank Name:",       _BANK["bank"]),
        ("A/c No.:",         _BANK["account"]),
        ("Branch & IFSC:",   _BANK["ifsc"]),
    ]:
        pdf.set_x(LM)
        pdf.set_font("Helvetica", "B", 7.5)
        pdf.cell(32, 4.5, f_lbl)
        pdf.set_font("Helvetica", "", 7.5)
        pdf.cell(bank_w - 32, 4.5, f_val, ln=True)
    bank_bot = pdf.get_y()

    sig_x = LM + bank_w + 2
    pdf.set_xy(sig_x, foot_y)
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(W - bank_w - 2, 5, f"For {_CO['name']}", align="R")
    pdf.set_xy(sig_x, foot_y + 20)
    pdf.cell(W - bank_w - 2, 5, "Authorised Signatory", align="R")

    pdf.set_draw_color(136, 136, 136)
    pdf.line(LM + bank_w, foot_y, LM + bank_w, max(bank_bot, foot_y + 25))

    return bytes(pdf.output())


# ── Main view ──────────────────────────────────────────────────────────────────

def render() -> None:
    st.markdown(
        "<div class='page-eyebrow'>// Finance</div>"
        "<div class='page-title'>Invoice</div>",
        unsafe_allow_html=True,
    )
    st.markdown("<div style='margin-top:24px'></div>", unsafe_allow_html=True)

    try:
        sb = SupabaseClient()
    except Exception as exc:
        st.error(f"Supabase connection failed: {exc}")
        return

    try:
        work_orders = sb.list_work_orders()
    except Exception as exc:
        st.error(f"Failed to load work orders: {exc}"); return

    customers = []
    sites     = []
    try: customers = sb.list_customers()
    except Exception: pass
    try: sites = sb.list_sites()
    except Exception: pass

    customer_map = {c["id"]: c for c in customers if c.get("id")}
    site_map     = {s["id"]: s for s in sites     if s.get("id")}
    wo_map       = {w["id"]: w for w in work_orders if w.get("id")}

    # ── Selectors: Customer → Site → Work Order ───────────────────────────────
    cids = sorted(
        {wo.get("customer_id") for wo in work_orders if wo.get("customer_id")},
        key=lambda c: customer_map.get(c, {}).get("customer_name", ""),
    )

    sel_cid = st.selectbox(
        "Customer",
        [""] + cids,
        format_func=lambda x: "Select customer" if not x
            else customer_map.get(x, {}).get("customer_name", x),
        key="inv_cid",
    )

    if st.session_state.get("_inv_prev_cid") != sel_cid:
        st.session_state["_inv_prev_cid"] = sel_cid
        st.session_state["inv_site"] = ""
        st.session_state["inv_wo"]   = ""

    if not sel_cid:
        st.markdown(
            "<div style='margin-top:32px;padding:40px;background:#f8fafc;"
            "border:1px dashed #d1d5db;border-radius:10px;text-align:center;'>"
            "<div style='font-size:32px;'>🧾</div>"
            "<div style='font-size:14px;font-weight:600;color:#374151;margin-top:10px;'>"
            "Select a Customer to begin.</div></div>",
            unsafe_allow_html=True,
        )
        return

    # Site — filtered to sites that have WOs for this customer
    _site_ids_for_cid = sorted(
        {wo.get("site_id") for wo in work_orders
         if wo.get("customer_id") == sel_cid and wo.get("site_id")},
        key=lambda sid: site_map.get(sid, {}).get("site_name", ""),
    )

    sc1, sc2 = st.columns(2)
    with sc1:
        sel_site_id = st.selectbox(
            "Site",
            [""] + _site_ids_for_cid,
            format_func=lambda x: "Select site" if not x
                else site_map.get(x, {}).get("site_name", x),
            key="inv_site",
            disabled=not sel_cid,
        )

    if st.session_state.get("_inv_prev_site") != sel_site_id:
        st.session_state["_inv_prev_site"] = sel_site_id
        st.session_state["inv_wo"] = ""

    # Work Order — filtered by customer + site
    wo_ids = sorted(
        [wid for wid, wo in wo_map.items()
         if wo.get("customer_id") == sel_cid and wo.get("site_id") == sel_site_id],
        key=lambda wid: wo_map[wid].get("wo_number", ""),
    ) if sel_site_id else []

    with sc2:
        sel_wo_id = st.selectbox(
            "Work Order",
            [""] + wo_ids,
            format_func=lambda x: "Select work order" if not x
                else wo_map[x].get("wo_number", "Unknown"),
            key="inv_wo",
            disabled=not sel_site_id,
        )

    if not sel_site_id or not sel_wo_id:
        st.markdown(
            "<div style='margin-top:32px;padding:40px;background:#f8fafc;"
            "border:1px dashed #d1d5db;border-radius:10px;text-align:center;'>"
            "<div style='font-size:32px;'>🧾</div>"
            "<div style='font-size:14px;font-weight:600;color:#374151;margin-top:10px;'>"
            "Select a Site and Work Order above to continue.</div></div>",
            unsafe_allow_html=True,
        )
        return

    sel_wo       = wo_map[sel_wo_id]
    sel_customer = customer_map.get(sel_wo.get("customer_id", ""), {})
    sel_site     = site_map.get(sel_wo.get("site_id", ""), {})

    raw_mc = sel_wo.get("machine_config")
    mc_list: list[dict] = []
    if raw_mc:
        try:
            recs = json.loads(raw_mc) if isinstance(raw_mc, str) else raw_mc
            mc_list = [r for r in (recs if isinstance(recs, list) else []) if r.get("machine_label")]
        except Exception:
            pass

    if not mc_list:
        st.warning("No machines configured on this work order.")
        return

    st.markdown("<div style='margin-top:6px'></div>", unsafe_allow_html=True)

    # ── Two-panel layout ───────────────────────────────────────────────────────
    left_col, right_col = st.columns([4, 7], gap="large")

    # ── LEFT: charge selection + config ───────────────────────────────────────
    with left_col:
        st.markdown(
            "<div style='font-size:10px;font-weight:700;letter-spacing:.12em;"
            "color:#E87722;text-transform:uppercase;margin-bottom:10px;'>"
            "Select Charges to Invoice</div>",
            unsafe_allow_html=True,
        )

        selected_items: list[dict] = []

        for i, mc in enumerate(mc_list):
            mid  = mc.get("machine_id") or str(i)
            mlbl = mc.get("machine_label", f"Machine {i+1}")

            with st.expander(f"🔧 {mlbl}", expanded=(i == 0)):
                # Completed worklogs — split into pending (not yet billed) and already billed
                all_wls: list[dict] = []
                try:
                    raw_wls = sb.list_worklogs_for_machine(sel_wo_id, mid)
                    all_wls = [w for w in raw_wls if not w.get("is_draft", True)]
                except Exception:
                    pass

                pending_wls = [w for w in all_wls if not w.get("invoiced")]
                billed_wls  = [w for w in all_wls if w.get("invoiced")]

                if pending_wls:
                    st.markdown(
                        "<div style='font-size:9px;font-weight:700;color:#6B7280;"
                        "text-transform:uppercase;margin-bottom:6px;'>Pending Billing</div>",
                        unsafe_allow_html=True,
                    )
                    for wl in pending_wls:
                        bm      = wl.get("year", "—")
                        billing = _compute_billing(wl, mc)
                        period  = _period_str(bm, mc)
                        wl_key  = f"inv_wl_{sel_wo_id}_{mid}_{bm}"
                        if st.checkbox(
                            f"{bm} — ₹ {_fmt_inr(billing['net'])}", key=wl_key, value=True
                        ):
                            selected_items.append({
                                "type": "worklog", "mc": mc, "wl": wl,
                                "billing": billing, "period": period, "sl": i + 1,
                            })
                elif not billed_wls:
                    st.caption("No completed worklogs for this machine.")

                if billed_wls:
                    st.markdown(
                        "<div style='font-size:9px;font-weight:700;color:#166534;"
                        "text-transform:uppercase;margin-bottom:4px;"
                        + ("margin-top:10px;" if pending_wls else "")
                        + "'>Already Billed</div>",
                        unsafe_allow_html=True,
                    )
                    for wl in billed_wls:
                        bm      = wl.get("year", "—")
                        inv_ref = wl.get("invoice_number") or "—"
                        billing = _compute_billing(wl, mc)
                        st.markdown(
                            f"<div style='padding:5px 10px;background:#f0fdf4;"
                            f"border:1px solid #86efac;border-radius:6px;font-size:11px;"
                            f"color:#166534;margin-bottom:4px;'>"
                            f"✓ <b>{bm}</b> — ₹ {_fmt_inr(billing['net'])}"
                            f"&nbsp;&nbsp;·&nbsp;&nbsp;Invoice: <b>{inv_ref}</b></div>",
                            unsafe_allow_html=True,
                        )

                mob  = float(mc.get("mobilization_cost") or 0)
                demob = float(mc.get("demobilization_cost") or 0)
                if mob > 0 and st.checkbox(
                    f"Mobilisation — ₹ {_fmt_inr(mob)}",
                    key=f"inv_mob_{sel_wo_id}_{mid}",
                ):
                    selected_items.append({"type": "mob", "mc": mc, "amount": mob, "sl": i + 1})

                if demob > 0 and st.checkbox(
                    f"Demobilisation — ₹ {_fmt_inr(demob)}",
                    key=f"inv_demob_{sel_wo_id}_{mid}",
                ):
                    selected_items.append({"type": "demob", "mc": mc, "amount": demob, "sl": i + 1})

                # Item code — auto-populated from the machine config (set in Work Orders)
                if st.session_state.get("inv_ic", False):
                    _mc_ic = mc.get("item_code") or ""
                    st.text_input(
                        "Item Code",
                        value=_mc_ic,
                        key=f"inv_ic_val_{mid}",
                        placeholder="e.g. SRV-001",
                        help="Auto-filled from Work Order machine config. Edit here to override for this invoice.",
                    )

        # ── Invoice config ─────────────────────────────────────────────────────
        st.markdown("<div style='margin-top:14px'></div>", unsafe_allow_html=True)
        st.markdown(
            "<div style='font-size:10px;font-weight:700;letter-spacing:.12em;"
            "color:#E87722;text-transform:uppercase;margin-bottom:10px;'>"
            "Invoice Configuration</div>",
            unsafe_allow_html=True,
        )

        prefix   = f"BL/CLI/{_fy_str()}/"
        inv_input = st.text_input(
            "Invoice Number",
            placeholder=f"{prefix}204",
            key="inv_number",
            help=f"Format: {prefix}NNN — enter just the number or the full invoice number",
        )
        inv_no = (
            inv_input.strip()
            if "/" in inv_input.strip()
            else f"{prefix}{inv_input.strip()}"
        ) if inv_input.strip() else ""

        inv_date  = st.date_input("Invoice Date", value=date.today(), key="inv_date")
        tax_on    = st.checkbox("Apply GST (18%)", value=True, key="inv_tax")
        tax_type  = "CGST/SGST"
        if tax_on:
            tax_type = st.radio(
                "Tax Type", ["CGST/SGST", "IGST"],
                key="inv_tax_type", horizontal=True,
                help="CGST/SGST = intra-state · IGST = inter-state",
            )
        hsn_on   = st.checkbox("Include HSN/SAC code", value=False, key="inv_hsn")
        hsn_code = st.text_input("HSN/SAC", placeholder="997319", key="inv_hsn_val") if hsn_on else ""
        ic_on    = st.checkbox("Include Item Codes", value=False, key="inv_ic")
        inv_format = st.radio(
            "Invoice Format",
            ["Option 1: With Company Letterhead", "Option 2: Blank Header (Pre-printed Letterhead)"],
            key="inv_format",
            help="Option 2 leaves the top header blank so you can print on your own pre-printed letterhead paper.",
        )
        blank_header = inv_format.startswith("Option 2")
        notes    = st.text_area("Notes (optional)", key="inv_notes", height=60)

        # duplicate check
        dup = False
        if inv_no:
            try:
                dup = sb.invoice_number_exists(inv_no)
                if dup:
                    st.error(f"Invoice number **{inv_no}** already exists.")
            except Exception:
                pass

        can_gen = bool(inv_no) and not dup and bool(selected_items)
        if not selected_items:
            st.warning("Select at least one charge above.")

        gen_btn = st.button(
            "🧾  Generate Invoice",
            type="primary",
            disabled=not can_gen,
            use_container_width=True,
            key="inv_gen",
        )

    # ── RIGHT: preview ─────────────────────────────────────────────────────────
    with right_col:
        st.markdown(
            "<div style='font-size:10px;font-weight:700;letter-spacing:.12em;"
            "color:#E87722;text-transform:uppercase;margin-bottom:10px;'>"
            "Invoice Preview</div>",
            unsafe_allow_html=True,
        )

        if not selected_items:
            st.markdown(
                "<div style='padding:60px;background:#f8fafc;border:1px dashed #d1d5db;"
                "border-radius:10px;text-align:center;'>"
                "<div style='font-size:40px;'>🧾</div>"
                "<div style='font-size:13px;color:#6B7280;margin-top:10px;'>"
                "Select charges on the left to preview the invoice.</div></div>",
                unsafe_allow_html=True,
            )
        else:
            # Build line groups
            grp_map: dict[str, dict] = {}
            sl_seen: dict[str, int]  = {}

            for item in selected_items:
                mc   = item["mc"]
                mid  = mc.get("machine_id") or str(item["sl"] - 1)
                mlbl = mc.get("machine_label", "")
                mtype = mlbl.split("—")[0].strip() if "—" in mlbl else mlbl

                if mid not in grp_map:
                    sl_seen[mid] = len(sl_seen) + 1
                    grp_map[mid] = {
                        "machine_label": mlbl,
                        "make":   mc.get("make", ""),
                        "model":  mc.get("model", ""),
                        "serial": mc.get("serial_number", ""),
                        "sl_no":  sl_seen[mid],
                        "item_code": st.session_state.get(f"inv_ic_val_{mid}", "") if ic_on else "",
                        "items": [],
                    }

                grp = grp_map[mid]

                if item["type"] == "worklog":
                    b = item["billing"]
                    p = item["period"]
                    grp["items"].append({
                        "desc":   f"Hiring charges - {mtype}\nPeriod - {p}",
                        "uom":    "Month",
                        "qty":    f"{b['qty']:.3f}",
                        "rate":   b["rate"],
                        "amount": b["hiring"],
                    })
                    if b["ot_hrs"] > 0:
                        grp["items"].append({
                            "desc":   f"OT charges - {mtype}\nPeriod - {p}",
                            "uom":    "Hourly",
                            "qty":    f"{b['ot_hrs']:.3f}",
                            "rate":   b["ot_rate"],
                            "amount": b["ot_amt"],
                        })
                    if b["deduction"] > 0:
                        grp["items"].append({
                            "desc":   "Deduction",
                            "uom":    "—",
                            "qty":    "—",
                            "rate":   0,
                            "amount": -b["deduction"],
                        })
                    if b["add_op_amt"] > 0:
                        grp["items"].append({
                            "desc":   "Additional Operator Charges",
                            "uom":    "Nos",
                            "qty":    f"{b['add_op_qty']:.0f}",
                            "rate":   b["add_op_rate"],
                            "amount": b["add_op_amt"],
                        })
                    if b["add_acc_amt"] > 0:
                        grp["items"].append({
                            "desc":   "Additional Accommodation Charges",
                            "uom":    "Nos",
                            "qty":    f"{b['add_acc_qty']:.0f}",
                            "rate":   b["add_acc_rate"],
                            "amount": b["add_acc_amt"],
                        })
                elif item["type"] == "mob":
                    grp["items"].append({
                        "desc":   "Mobilisation charges",
                        "uom":    "Nos",
                        "qty":    "1.000",
                        "rate":   item["amount"],
                        "amount": item["amount"],
                    })
                elif item["type"] == "demob":
                    grp["items"].append({
                        "desc":   "Demobilisation charges",
                        "uom":    "Nos",
                        "qty":    "1.000",
                        "rate":   item["amount"],
                        "amount": item["amount"],
                    })

            groups = list(grp_map.values())

            inv_html = _build_html(
                inv_no=inv_no or f"{prefix}???",
                inv_date=inv_date,
                wo=sel_wo,
                customer=sel_customer,
                site=sel_site,
                groups=groups,
                tax_type=tax_type,
                tax_on=tax_on,
                hsn_on=hsn_on,
                hsn_code=hsn_code,
                item_code_on=ic_on,
                notes=notes.strip() if notes else "",
                blank_header=blank_header,
            )

            st.iframe(inv_html, height=960)

            # ── Action row ─────────────────────────────────────────────────────
            _inv_fname = inv_no.replace("/", "_") if inv_no else "invoice"
            _docx_kwargs = dict(
                inv_no=inv_no or f"{prefix}???",
                inv_date=inv_date,
                wo=sel_wo, customer=sel_customer, site=sel_site,
                groups=groups, tax_type=tax_type, tax_on=tax_on,
                hsn_on=hsn_on, hsn_code=hsn_code,
                item_code_on=ic_on,
                notes=notes.strip() if notes else "",
                blank_header=blank_header,
            )
            # Build file bytes once — shared by buttons and storage upload
            _docx_bytes: bytes | None = None
            _pdf_bytes:  bytes | None = None
            try:
                _docx_bytes = _build_docx(**_docx_kwargs)
            except Exception:
                pass
            try:
                _pdf_bytes = _build_pdf_bytes(**_docx_kwargs)
            except Exception:
                pass

            a1, a2, a3, a4, a5 = st.columns(5)
            with a1:
                st.download_button(
                    "⬇  Download (HTML)",
                    data=inv_html.encode("utf-8"),
                    file_name=f"{_inv_fname}.html",
                    mime="text/html",
                    use_container_width=True,
                )
            with a2:
                if _docx_bytes:
                    st.download_button(
                        "⬇  Download (Word)",
                        data=_docx_bytes,
                        file_name=f"{_inv_fname}.docx",
                        mime="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document",
                        use_container_width=True,
                    )
                else:
                    st.error("Word export failed")
            with a3:
                if _pdf_bytes:
                    st.download_button(
                        "⬇  Download (PDF)",
                        data=_pdf_bytes,
                        file_name=f"{_inv_fname}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                    )
                else:
                    st.error("PDF export failed")
            with a4:
                if st.button(
                    "🖨️  Print + Work Log",
                    key="btn_print_with_wl",
                    use_container_width=True,
                ):
                    st.session_state["_inv_print_mode"] = "with_wl"
            with a5:
                if gen_btn and can_gen:
                    subtotal = sum(sum(it["amount"] for it in g["items"]) for g in groups)
                    tax_tot  = round(subtotal * 0.18, 2) if tax_on else 0.0
                    grand    = round(subtotal + tax_tot)
                    rnd_off  = round(grand - (subtotal + tax_tot), 2)
                    try:
                        sb.insert_invoice({
                            "invoice_number": inv_no,
                            "work_order_id":  sel_wo_id,
                            "invoice_date":   inv_date.isoformat(),
                            "customer_id":    sel_wo.get("customer_id"),
                            "site_id":        sel_wo.get("site_id"),
                            "tax_type":       tax_type,
                            "line_items":     json.dumps(groups),
                            "subtotal":       subtotal,
                            "tax_amount":     tax_tot,
                            "round_off":      rnd_off,
                            "grand_total":    grand,
                            "status":         "Final",
                            "notes":          notes.strip() or None,
                        })
                        # Mark every worklog line-item as invoiced so they
                        # disappear from the Pending for Billing list.
                        for _item in selected_items:
                            if _item.get("type") == "worklog":
                                _wl_id = (_item.get("wl") or {}).get("id")
                                if _wl_id:
                                    try:
                                        sb.mark_worklog_invoiced(_wl_id, inv_no)
                                    except Exception:
                                        pass
                        # Upload Word + PDF to storage for future downloads
                        _upload_ok: list[str] = []
                        for _ext, _fbytes in [("docx", _docx_bytes), ("pdf", _pdf_bytes)]:
                            if _fbytes:
                                try:
                                    sb.upload_invoice_file(inv_no, _fbytes, _ext)
                                    _upload_ok.append(_ext.upper())
                                except Exception:
                                    pass
                        _saved_msg = f"✔ Invoice **{inv_no}** saved."
                        if _upload_ok:
                            _saved_msg += f" Files stored: {', '.join(_upload_ok)}."
                        st.success(_saved_msg)
                    except Exception as exc:
                        st.warning(f"Preview ready — could not save to DB: {exc}")

            # ── Print Invoice + Work Log preview ───────────────────────────────
            if st.session_state.get("_inv_print_mode") == "with_wl":
                wl_section = _build_worklog_schedules_html(selected_items)
                if wl_section:
                    combined_html = inv_html.replace(
                        "</body></html>",
                        f"{wl_section}\n</body></html>",
                    )
                    st.markdown(
                        "<div style='margin-top:12px;font-size:10px;font-weight:700;"
                        "letter-spacing:.1em;color:#E87722;text-transform:uppercase;"
                        "margin-bottom:6px;'>Invoice + Work Log (Print Preview)</div>",
                        unsafe_allow_html=True,
                    )
                    st.iframe(combined_html, height=1200)
                else:
                    st.info("No completed work log schedules found for the selected charges.")

    # ── Documents (attached to the work order's invoice record) ───────────────
    with st.expander("📎 Documents", expanded=False):
        render_document_panel(
            sb,
            record_type = "invoice",
            record_id   = sel_wo_id,
            key_prefix  = "inv",
        )
